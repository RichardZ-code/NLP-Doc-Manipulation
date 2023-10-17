# Import the libraries
import docx
import os.path
from docx import Document
from docx.shared import Inches


# Prepare to read the Word document named 'manuscript.docx'
docxFile = 'manuscript.docx'
document = docx.Document(docxFile)


# Create a Word document to store the figure legends temporarily
figure_legends = Document()


# To determine if [para] is followed by a figure to be inserted
def if_contains(figure_name: str, para):

    if figure_name in para.text:
        return True
    else:
        return False


# Define the figure legend index and the figure index, both started from 1
figure_legend_index = 1
figure_index = 1


# Define the paragraph index, started from 0
paragraph_index = 0


# Loop from the first paragraph to the last paragraph in document
for i in document.paragraphs:

    # Define the figure legend name
    figure_legend_name = 'Figure' + ' ' + str(figure_legend_index) + '.'

    # Add i to 'figure_legends.docx' and remove i from 'manuscript.docx' if i is a figure legend
    if if_contains(figure_legend_name, i):
        paragraph = figure_legends.add_paragraph(i.text)
        i.clear()
        figure_legend_index += 1


# Scan the whole document again
for i in document.paragraphs:

    # Define the names
    figure_name = 'Figure' + ' ' + str(figure_index)
    figures_name = 'Figures' + ' ' + str(figure_index)
    figure_name_tif = 'Figure' + ' ' + str(figure_index) + '.tif'

    # Insert the corresponding figure and figure legend if i is the right paragraph
    if if_contains(figure_name, i) or if_contains(figures_name, i):
        figure_paragraph = document.paragraphs[paragraph_index + 1].insert_paragraph_before()
        run = figure_paragraph.add_run()
        run.add_picture(figure_name_tif)
        run.add_text(figure_legends.paragraphs[figure_index - 1].text)
        figure_index += 1
        # Add one to the paragraph index since a figure paragraph was inserted
        paragraph_index += 1

    # The paragraph index is incremented by one
    paragraph_index += 1


# Save the document as 'manuscript.docx'
document.save('manuscript.docx')