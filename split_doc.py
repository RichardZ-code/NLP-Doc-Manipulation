# To import the libraries
import glob
from docx import Document


# Define the split_doc index, started from 0
index = 0


def is_heading(para):
    """
    This function determines if [para] is a heading
    """
    if 'LIHC_' in para.text:
        return True
    else:
        return False


def iterate_document_sections(doc):
    """
    This function generates a sequence of paragraphs for each headed section, 
    where each sequence starts with a heading and followed by its contents
    """
    paragraphs = [doc.paragraphs[0]]
    for paragraph in doc.paragraphs[1:]:
        if is_heading(paragraph):
            yield paragraphs
            paragraphs = [paragraph]
            continue
        paragraphs.append(paragraph)
    yield paragraphs


def create_document_from_paragraphs(paragraphs):
    """
    This function creates a Word document for the headed section and
    saves it as 'split_docX.docx'
    """

    # To make 'index' global
    global index

    split_doc = Document()
    for counter, words in enumerate(paragraphs):
        new_content = words.text
        split_doc.add_paragraph(new_content)
        print(new_content)
    split_doc.save('split_doc' + str(index) + '.docx')

    # To name the Word documents in a numerical sequence: 0, 1, 2 ...
    index += 1


# Find the file named '41_50.docx'
for name in glob.glob('41_50.docx'):
    document = Document(name)
    for paragraphs in iterate_document_sections(document):
        create_document_from_paragraphs(paragraphs)